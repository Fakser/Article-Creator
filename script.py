from nltk.corpus import words
from nltk.corpus import wordnet
from nltk.corpus import stopwords
from nltk.metrics import edit_distance 
from time import sleep
from time import time
from bs4 import BeautifulSoup
from googlesearch import search
from copy import deepcopy
from docx import Document
from PIL import Image
from numpy.random import uniform
from sys import argv

import string
import shutil
import os
import urllib
import requests
import re
import language_check
import random

# NLP and webscrapping based article/report creator
# author: Krzysztof Kramarz

# our hyperparameters
topic = 'machine learning basics' # Topic of an article
sentence_quality = 0.4 # variable describing quality of the sentence
number_of_articles = 10 # number of articles scipt will be scrapping
change_to_synonym_chance = 0.2
article_randomness = 0.2
max_word_len = 10
tool = language_check.LanguageTool('en-US') # tool for checking grammar in sentences
delete_unnecesarry = True

if len(argv) > 1:
    topic = argv[1]

start = time()

# getting urls of pages that will be scraped
# first we will use googlesearch module for getting results from google
google_search = search('everything about ' + topic)
urls = []
number_of_proper_articles = 0

print('getting urls')

# we will go through these results as long as we will find enough different pages 
while number_of_proper_articles < number_of_articles:
    # we get next url from the generator
    new_url = next(google_search)
    bad_url = False
    # next we check if these page was not already added to the list (google likes to give you the same result many times)
    for url in urls:
        if url in new_url or new_url in url or 'youtube' in new_url or 1/(edit_distance(new_url, url)+1) > 0.05 or uniform() < article_randomness:
            bad_url = True
            break
    if bad_url == False:
        urls.append(new_url)
        number_of_proper_articles += 1
        print(new_url)
print('_________________________________________')        

# next we will get response html with requests.get(url) and split it into text and images
articles = []
images = []
for url in urls:
    # downloading response 
    try:
        page = requests.get(url)
        print(url + '    Downloaded')
    except:
        print(url + '    Failed')
        continue
    # splitting response by bracets to get code part and text part separetly - i found this way the fastest and most proficient
    soup = BeautifulSoup(page.content, 'html.parser')
    page_splitted = re.split('{|}',soup.text)
    for article in page_splitted:
        # cleaning text part from unnecesary stuff
        article = article.replace('\n', ' ').replace('  ','').replace('\\u','')
        new_article = ''
        for sentence in article.split('. '):
            new_article += re.sub('["#$%&()*+-/:;<=>@^_{|}~\r\t]', '', sentence) + '. '
        # if this text meets our requierments we add it to the list of articles
        if len(new_article) >= 4000:
            articles.append(re.sub('(figure|fig|fig.) [0-9]', '' ,new_article))
    # iterating throught all img objects found in responce to get some hot pictures
    for index, image in enumerate(soup.findAll('img')):
        if 'srcset' in image.attrs.keys():
            new_url = image.attrs['srcset'].split(',')[0]
            bad_url = False
            for url in images:
                if (1/edit_distance(new_url, url)+1) > 0.2:
                    bad_url = True
                    break
            if bad_url == False:
                images.append(new_url)
                print(new_url)
        # if index > int(len(soup.findAll('img')) * 4/7) and index < int(len(soup.findAll('img')) * 5/7) and 'src' in image.attrs.keys():
        #     images.append(image.attrs['src'])

    sleep(0.2)

print('getting images')
# downloading photos from the list of image urls
# creating directory for photos
os.mkdir('./photos')
for index, image_url in enumerate(images):
    # getting path for file
    if '.png' in image_url:
        file_type = '.png'
    else:
        file_type = '.jpg'
    file_path = './photos/' + str(index) + 'bonus' + file_type
    delete = False
    # creating file from url with many conditions 
    with open(file_path, 'wb') as photo:
        try:
            photo.write(requests.get(image_url.split(' ')[0]).content)
        except:
            try:
                photo.write(request.urlopen(image_url).read())
            except:
                delete = True
                continue
    try:
        # checking if image is not corrupted (not really working yet) and resizing it so it can fit the document
        im = Image.open(file_path)
        im.verify()
        im.close()
        basewidth = 500
        img = Image.open(file_path)
        if img.size[0] > 500:
            wpercent = (basewidth/float(img.size[0]))
            hsize = int((float(img.size[1])*float(wpercent)))
            img = img.resize((basewidth,hsize), Image.ANTIALIAS)
        img.save(file_path)
        img.close()
    except Exception as ex:
        try:
            os.remove('./photos/' + str(index) + 'bonus' + file_type)
        except:
            continue
        print(ex) 
    if delete:
        try:
            os.remove('./photos/' + str(index) + 'bonus' + file_type)
        except:
            continue
    
def get_synonyms_and_antonyms(word):
    # function that returns list of synonyms ant antonyms for given word
    synonyms = []
    antonyms = []

    for syn in wordnet.synsets(wprd):
        for l in syn.lemmas():
            synonyms.append(l.name())
            if l.antonyms():
                    antonyms.append(l.antonyms()[0].name())
    return synonyms, antonyms

print('cleaning sentences')

# quality of sentence metrics
# we need to get only most sensical senetences with some gramarr cleaning and word replacements
english_dictonary = set(words.words())
cleaned_articles = [] 
for article in articles:
    cleaned_article = []
    # splitting article into setentces
    for sentence in article.split('. '):
        sentence_to_words = re.split(' ', sentence)
        # measure of sentence sense and cleaning it
        metrics = 0
        new_sentence = ''
        for word in sentence_to_words:
            if word in english_dictonary or len(word) <max_word_len:
                metrics += 1
                # adding word or its synonym to the sentence
                if uniform() < change_to_synonym_chance:
                    try:
                        synonyms, _ = get_synonyms_and_antonyms(word)
                        new_sentence +=  random.choice(synonyms)  
                    except:
                        new_sentence += word + ' '
                else:
                    new_sentence += word + ' '

        # adding sentence if it meets requeirments
        metrics = metrics/len(sentence_to_words)
        new_sentence = new_sentence[:-1]
        try:
            matches = tool.check(new_sentence)
            for match in matches:
                if 'anise' in match.replacements or 'mange' in match.replacements:
                    matches.remove(match)
            new_sentence = language_check.correct(new_sentence, matches)
            if metrics >= sentence_quality:
                cleaned_article.append(new_sentence)
        except:
            continue
    cleaned_articles.append(deepcopy(cleaned_article))

print('creating random article')

# additional clearing from unnecessary parts like terms of use etc.
if delete_unnecesarry == True:
    for index in range(len(cleaned_articles)):
        cleaned_articles[index] = cleaned_articles[index][int(len(cleaned_articles[index])* 0.1):int(len(cleaned_articles[index])* 0.9)]

# random new article
# splitting each article into introduction, elaboration and conclusion and mixing them 
random_article = {'Introduction': [], 'Elaboration': [], 'Conclusions': []}

for article in cleaned_articles:
    if len(article) > 10:
        introduction = article[:int(len(article)/6)]
        elaboration = article[int(len(article)/6):int(len(article)*5/6)]
        conclusions = article[int(len(article)*5/6):]
        split_1 = random.randint(1, len(introduction))
        split_2 = random.randint(1, len(elaboration))
        random_article['Introduction'] += introduction[int(split_1/2) : split_1]
        random_article['Elaboration'] += elaboration[int(split_2/2): split_2]
        random_article['Conclusions'] += conclusions[int(split_1/2) : split_1]
    else:
        continue

# some data science - if you would like to know dictonary of your article
word_count = 0
word_dict = {}
for key in random_article.keys():
    for sentence in random_article[key]:
        word_count += len(re.split(' |\n',  sentence))
        for word in re.split(' |\n',  sentence):
            if word in word_dict.keys():
                word_dict[word] += 1
            else:
                word_dict[word] = 1 

def get_header(text, lenght):
    # function that creates header for paragraph 
    # deleting stopwords
    text = ' '.join([word for word in text.lower().split() if word not in (stopwords.words('english'))])
    # creating dictionary with counts for each word
    dictonary = {}
    for word in re.split('[ ?!,.]', text):
        if len(word) > 4:
            new_word = True
            for key in dictonary.keys():
                if word == key:
                    new_word = False
                    dictonary[word] += 1
                    break
                if word in key:
                    new_word = False
                    dictonary[word] = dictonary.pop(key)
                    break
                if key in word:
                    new_word = False
                    break
            if new_word:
                dictonary[word] = 1
    # first we add most frequent words        
    header = ''
    for index, word in enumerate(sorted(dictonary, key=dictonary.get, reverse=True)):
        header += word
        if index ==  lenght:
            header = header[:1].upper() + header[1:]
            # next we try to create a sensical sentence from it
            matches = tool.check(header)
            for match in matches:
                if 'anise' in match.replacements or 'mange' in match.replacements:
                    matches.remove(match)
            return language_check.correct(header, matches)
            
        header += ' '

# creation of .docx
print('creating and saving docx file')

document = Document()
document.add_heading(topic.upper(), 0)

m = len(os.listdir('./photos/'))
n = sum([len(random_article[key]) for key in random_article.keys()])
if n > m:
    r = m/(n * 2)
else:
    r = n/(m * 2)
image_chance = 1

for key in random_article.keys():
    split_index = random.randint(10,20)
    part = ''
    for index, sentence in enumerate(random_article[key]):
        part += sentence + '. '
        if index == split_index:
            document.add_heading(get_header(part, random.randint(1,3)), 2)
            paragraph = document.add_paragraph()
            paragraph.add_run(part)
            split_index += random.randint(10,15)
            part = ''
        image_chance -=  r
        if uniform() > image_chance:
            try:
                filename = os.listdir('./photos/')[0]
                try:
                    document.add_picture('./photos/' + filename)
                except:
                    pass
                image_chance = 1
                os.remove('./photos/' + filename)
            except Exception as ex:
                continue

            
    document.add_heading(get_header(part, random.randint(2,4)), 2)
    paragraph = document.add_paragraph()
    paragraph.add_run(part)
document.save('report.docx')
shutil.rmtree('./photos')

print('whole process took: ', (time()  - start)/60, 'min')